
# from fastapi import FastAPI, File, UploadFile, Form, HTTPException
# from fastapi.middleware.cors import CORSMiddleware
# import fitz
# import pdfplumber
# import pytesseract
# from PIL import Image
# import io
# import docx
# import json
# import uvicorn
# from openai import OpenAI
# import re
# from datetime import datetime
# from pydantic import BaseModel
# from dotenv import load_dotenv
# import os 

# app = FastAPI()

# load_dotenv()
# api_key = os.getenv("OPENAI_API_KEY")
# openai = OpenAI(api_key=api_key)


# # ✅ הוספת CORS
# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],  # ניתן לשנות לכתובת ספציפית ["http://localhost:5173"]
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# # פונקציות לשליפת טקסט
# def extract_text_from_pdf(file_stream):
#     """ חילוץ טקסט מקובץ PDF """
#     text = ""
#     try:
#         doc = fitz.open(stream=file_stream, filetype="pdf")
#         for page in doc:
#             text += page.get_text()
#         return text.strip()
#     except Exception as e:
#         return f"Error extracting text: {str(e)}"

# def extract_text_from_image(file):
#     image = Image.open(io.BytesIO(file))
#     text = pytesseract.image_to_string(image, lang="heb+eng")
#     return text.strip()

# def extract_text_from_docx(file):
#     doc = docx.Document(io.BytesIO(file))
#     text = "\n".join([para.text for para in doc.paragraphs])
#     return text.strip()

# # פונקציה לחילוץ תאריך מתוך טקסטי
# def extract_date_from_text(text):
#     """ פונקציה לחילוץ תאריך מתוך טקסט """
#     date_pattern = r"(\d{1,2})[./-](\d{1,2})[./-](\d{4})"
#     matches = re.findall(date_pattern, text)
#     print("Matches found:", matches)  # הדפסת התאריכים שנמצאו

#     if matches:
#         for match in matches:
#             day, month, year = match
#             # נוודא שהתאריך תקין
#             try:
#                 date_str = f"{year}-{month.zfill(2)}-{day.zfill(2)}"
#                 return datetime.strptime(date_str, "%Y-%m-%d").date()
#             except ValueError:
#                 continue
#     return None

# @app.get("/")
# def root():
#     return {"status": "Server running"}

# @app.post("/upload/")
# async def upload_file(file: UploadFile = File(...), event_type: str = Form(...)):
#     file_bytes = await file.read()

#     # בדיקה אם הקובץ ריק
#     if not file_bytes:
#         raise HTTPException(status_code=400, detail="הקובץ ריק")

#     extracted_text = ""
    
#     # בדיקה וסינון סוג הקובץ
#     if file.content_type == "application/pdf":
#         extracted_text = extract_text_from_pdf(io.BytesIO(file_bytes))
#         print(extracted_text)
#     elif file.content_type.startswith("image/"):
#         extracted_text = extract_text_from_image(file_bytes)
#     elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#         extracted_text = extract_text_from_docx(file_bytes)
#     else:
#         raise HTTPException(status_code=400, detail="סוג הקובץ אינו נתמך")

#     if not extracted_text:
#         raise HTTPException(status_code=400, detail="לא נמצא טקסט בקובץ")

#     # חילוץ תאריך מהטקסט
#     extracted_date = extract_date_from_text(extracted_text)
#     print(extracted_date)

#     if event_type == "חתונה":
#         extraction_prompt = """
#         You are an AI that extracts wedding invitation details.  
#         Always return **ONLY** a valid JSON object with the following structure:  
#         {
#             "שם כלה": "שם החתן",
#             "שם חתן": "שם הכלה",
#             "תאריך": "DD-MM-YYYY",
#             "שעה": "שעה",
#             "מקום": "מקום"
#         }
#         If a field is missing, set its value to null.  
#         Respond **only** with the JSON object (nothing else).  
#         """
#     else:
#         extraction_prompt = """
#         You are an AI that extracts birthday invitation details.  
#         Always return **ONLY** a valid JSON object with the following structure:  
#         {
#             "שם": "שם ",
#             "תאריך": "DD-MM-YYYY",
#             "שעה": "שעה",
#             "מקום": "מקום",
#             "סוג": "סוג",
#             "גיל": "גיל",
#         }
#         If a field is missing, set its value to null.  
#         Respond **only** with the JSON object (nothing else).  
#         """

#     if not extraction_prompt:
#         raise HTTPException(status_code=400, detail="סוג האירוע אינו תקין")

#     try:
#         response = openai.chat.completions.create(
#             model="gpt-4-turbo",
#             messages=[
#                 {"role": "system", "content": extraction_prompt},
#                 {"role": "user", "content": extracted_text}
#             ]
#         )
#         json_response = json.loads(response.choices[0].message.content)
        
#         # הוספת התאריך שהופק מהטקסט (אם נמצא)
#         # if extracted_date:
#         #     json_response["wedding_date"] = str(extracted_date)  # תאריך החתונה
#         # else:
#         #     json_response["wedding_date"] = None  # אם לא נמצא תאריך

#     except Exception as e:
#         raise HTTPException(status_code=500, detail=f"שגיאה בעיבוד הנתונים: {str(e)}")

#     return json_response

# @app.post("/marketing/generate")
# async def generate_marketing_text():
#     try:
#         prompt = f"""
#             כתוב מייל שיווקי בעברית תקנית, קלילה ומזמינה, שמטרתו להכיר לנמען את InvitationLine – אתר חדש לעיצוב הזמנות דיגיטליות לאירועים.

#             המייל צריך לכלול:
#             - כותרת מושכת שמפעילה רגש (שימחה, התרגשות, חיסכון בזמן, הקלה בתכנון).
#             - פנייה בגוף שני (את/אתם), בגובה העיניים, בגישה חמה ואישית.
#             - פתיחה שמציגה את הרעיון של InvitationLine והחזון שמאחוריו – לעשות עיצוב הזמנה פשוט, מהיר, יפה ומרגש.
#             - תיאור קצר של היתרונות: מגוון עיצובים ייחודיים, עורך נוח לשימוש, שליחה מיידית, התאמה אישית מלאה, ללא צורך בידע טכני.
#             - ניסוח משכנע אך לא מכירתי מדי – בלי קלישאות שיווקיות ריקות, אלא מסרים שנשמעים כמו המלצה אמיתית.
#             - קריאה ברורה לפעולה: "היכנסו עכשיו", "עצבו בעצמכם", או "נסו אותנו בחינם".
#             - חתימה בשם InvitationLine.

#             סגנון כתיבה: מקצועי אבל בגובה העיניים, נעים, חם, כזה שיגרום לקורא להרגיש שאנחנו מבינים אותו.
#             """
#         # f"""
#         #     כתוב מייל שיווקי בעברית תקנית בנושא עיצוב הזמנות לאירועים. 
#         #     המייל צריך לכלול:
#         #     - כותרת שיווקית מושכת.
#         #     - פנייה אישית ומזמינה.
#         #     - תוכן שיווקי ברור, מנוסח היטב, עם פסקאות מסודרות.
#         #     - משפט קריאה לפעולה בסיום.
#         #     לצוות שלנו קוראים invitationline , 
#         #     אתה יכול לחתום עם השם הזה.
#         #     """

#         response = openai.chat.completions.create(
#             model="gpt-4",
#             messages=[
#                 {"role": "system", "content": "אתה כותב תוכן שיווקי בעברית בצורה מקצועית."},
#                 {"role": "user", "content": f"כתוב מייל שיווקי בנושא: {prompt}"}
#             ],
#             max_tokens=1500
#         )

#         generated_text = response.choices[0].message.content.strip()
#         return {"text": generated_text}

#     except Exception as e:
#         return {"error": str(e)}

# if __name__ == "__main__":
#     # uvicorn.run("server:app", host="0.0.0.0", port=5000, reload=True)
#     port = int(os.environ.get("PORT", 5000))  # ברירת מחדל ל־5000 להרצה מקומית
#     uvicorn.run("server:app", host="0.0.0.0", port=port)




from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
import fitz
import pdfplumber
import pytesseract
from PIL import Image
import io
import docx
import json
import uvicorn
from openai import OpenAI
import re
from datetime import datetime
from pydantic import BaseModel
from dotenv import load_dotenv
import os 

app = FastAPI()

load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")
openai = OpenAI(api_key=api_key)


# ✅ הוספת CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # ניתן לשנות לכתובת ספציפית ["http://localhost:5173"]
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# פונקציות לשליפת טקסט
def extract_text_from_pdf(file_stream):
    """ חילוץ טקסט מקובץ PDF """
    text = ""
    try:
        doc = fitz.open(stream=file_stream, filetype="pdf")
        for page in doc:
            text += page.get_text()
        return text.strip()
    except Exception as e:
        return f"Error extracting text: {str(e)}"

def extract_text_from_image(file):
    image = Image.open(io.BytesIO(file))
    text = pytesseract.image_to_string(image, lang="heb+eng")
    return text.strip()

def extract_text_from_docx(file):
    doc = docx.Document(io.BytesIO(file))
    text = "\n".join([para.text for para in doc.paragraphs])
    return text.strip()

# פונקציה לחילוץ תאריך מתוך טקסטי
def extract_date_from_text(text):
    """ פונקציה לחילוץ תאריך מתוך טקסט """
    date_pattern = r"(\d{1,2})[./-](\d{1,2})[./-](\d{4})"
    matches = re.findall(date_pattern, text)
    # print("Matches found:", matches)  # הדפסת התאריכים שנמצאו

    if matches:
        for match in matches:
            day, month, year = match
            # נוודא שהתאריך תקין
            try:
                date_str = f"{year}-{month.zfill(2)}-{day.zfill(2)}"
                return datetime.strptime(date_str, "%Y-%m-%d").date()
            except ValueError:
                continue
    return None

@app.get("/")
def root():
    return {"status": "Server running"}

@app.post("/upload/")
async def upload_file(file: UploadFile = File(...), event_type: str = Form(...)):
    file_bytes = await file.read()

    # בדיקה אם הקובץ ריק
    if not file_bytes:
        raise HTTPException(status_code=400, detail="הקובץ ריק")

    extracted_text = ""
    
    # בדיקה וסינון סוג הקובץ
    if file.content_type == "application/pdf":
        extracted_text = extract_text_from_pdf(io.BytesIO(file_bytes))
        # print(extracted_text)
    elif file.content_type.startswith("image/"):
        extracted_text = extract_text_from_image(file_bytes)
    elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        extracted_text = extract_text_from_docx(file_bytes)
    else:
        raise HTTPException(status_code=400, detail="סוג הקובץ אינו נתמך")

    if not extracted_text:
        raise HTTPException(status_code=400, detail="לא נמצא טקסט בקובץ")

    # חילוץ תאריך מהטקסט (לא חובה להוסיף לשימוש ב-AI בלבד)
    extracted_date = extract_date_from_text(extracted_text)
    # print(extracted_date)

    if event_type == "חתונה":
        extraction_prompt = """
        You are an AI that extracts wedding invitation details from Hebrew and English text.
        You must identify and extract any date mentioned, including Hebrew calendar dates written in Hebrew words like "כ"ז סיון" or "ט' אדר".
        Do NOT convert Hebrew dates to the Gregorian calendar; return them exactly as they appear in the text.
        Return a valid JSON object only with this structure:
        {
            "שם כלה": "שם הכלה",
            "שם חתן": "שם החתן",
            "תאריך": "the date exactly as it appears in the text",
            "שעה": "שעה",
            "מקום": "מקום"
        }
        If a field is missing, set it to null.
        Respond only with the JSON object, no explanations or extra text.
        """
    else:
        extraction_prompt = """
        You are an AI that extracts birthday invitation details from Hebrew and English text.
        You must identify and extract any date mentioned, including Hebrew calendar dates written in Hebrew words like "כ"ז סיון" or "ט' אדר".
        Do NOT convert Hebrew dates to the Gregorian calendar; return them exactly as they appear in the text.
        Return a valid JSON object only with this structure:
        {
            "שם": "שם",
            "תאריך": "the date exactly as it appears in the text",
            "שעה": "שעה",
            "מקום": "מקום",
            "סוג": "סוג",
            "גיל": "גיל"
        }
        If a field is missing, set it to null.
        Respond only with the JSON object, no explanations or extra text.
        """

    if not extraction_prompt:
        raise HTTPException(status_code=400, detail="סוג האירוע אינו תקין")

    try:
        response = openai.chat.completions.create(
            model="gpt-4-turbo",
            messages=[
                {"role": "system", "content": extraction_prompt},
                {"role": "user", "content": extracted_text}
            ]
        )
        json_response = json.loads(response.choices[0].message.content)

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"שגיאה בעיבוד הנתונים: {str(e)}")

    return json_response

@app.post("/marketing/generate")
async def generate_marketing_text():
    try:
        prompt = f"""
            כתוב מייל שיווקי בעברית תקנית, קלילה ומזמינה, שמטרתו להכיר לנמען את InvitationLine – אתר חדש לעיצוב הזמנות דיגיטליות לאירועים.

            המייל צריך לכלול:
            - כותרת מושכת שמפעילה רגש (שימחה, התרגשות, חיסכון בזמן, הקלה בתכנון).
            - פנייה בגוף שני (את/אתם), בגובה העיניים, בגישה חמה ואישית.
            - פתיחה שמציגה את הרעיון של InvitationLine והחזון שמאחוריו – לעשות עיצוב הזמנה פשוט, מהיר, יפה ומרגש.
            - תיאור קצר של היתרונות: מגוון עיצובים ייחודיים, עורך נוח לשימוש, שליחה מיידית, התאמה אישית מלאה, ללא צורך בידע טכני.
            - ניסוח משכנע אך לא מכירתי מדי – בלי קלישאות שיווקיות ריקות, אלא מסרים שנשמעים כמו המלצה אמיתית.
            - קריאה ברורה לפעולה: "היכנסו עכשיו", "עצבו בעצמכם", או "נסו אותנו בחינם".
            - חתימה בשם InvitationLine.

            סגנון כתיבה: מקצועי אבל בגובה העיניים, נעים, חם, כזה שיגרום לקורא להרגיש שאנחנו מבינים אותו.
            """

        response = openai.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "אתה כותב תוכן שיווקי בעברית בצורה מקצועית."},
                {"role": "user", "content": f"כתוב מייל שיווקי בנושא: {prompt}"}
            ],
            max_tokens=1500
        )

        generated_text = response.choices[0].message.content.strip()
        return {"text": generated_text}

    except Exception as e:
        return {"error": str(e)}

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))  # ברירת מחדל ל־5000 להרצה מקומית
    uvicorn.run("server:app", host="0.0.0.0", port=port)
